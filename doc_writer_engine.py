#!/usr/bin/env python3
"""Artome Audio Document Writer - Multi-Format Exporter & Sharing Engine."""

import os
import shutil
import markdown
from docx import Document
from fpdf import FPDF

class AudioDocWriter:
    """Voice document writer with export support for PDF, DOCX, TXT, HTML, MD, and Cloud Sharing."""

    def __init__(self, output_dir=None):
        self.output_dir = output_dir or os.path.expanduser("~/Documents/ArtomeDocs")
        os.makedirs(self.output_dir, exist_ok=True)
        self.doc_title = "Untitled Document"
        self.sections = []  # List of {"type": "heading"|"paragraph", "content": str}

    def start_new_doc(self, title="Untitled Document"):
        """Initialize a new empty voice document."""
        self.doc_title = title.strip()
        self.sections = []
        return f"Started new document titled '{self.doc_title}'."

    def add_heading(self, heading_text):
        """Add a heading section to the document."""
        self.sections.append({"type": "heading", "content": heading_text.strip()})
        return f"Added heading: '{heading_text}'."

    def add_paragraph(self, paragraph_text):
        """Add a text paragraph to the document."""
        self.sections.append({"type": "paragraph", "content": paragraph_text.strip()})
        return f"Added paragraph: '{paragraph_text[:40]}...'"

    def read_draft_audio(self):
        """Read the complete document draft out loud."""
        if not self.sections:
            return f"Document '{self.doc_title}' is currently empty."

        speech = f"Reading document '{self.doc_title}'. Contains {len(self.sections)} sections. "
        for idx, sec in enumerate(self.sections, 1):
            if sec['type'] == 'heading':
                speech += f"Heading: {sec['content']}. "
            else:
                speech += f"Paragraph: {sec['content']}. "
        return speech

    def export_all_formats(self, filename_base=None):
        """Export document to TXT, Markdown, HTML, DOCX, and PDF simultaneously."""
        if not filename_base:
            filename_base = self.doc_title.lower().replace(" ", "_")

        base_path = os.path.join(self.output_dir, filename_base)
        exported_files = []

        # 1. Plain Text (.txt)
        txt_path = f"{base_path}.txt"
        with open(txt_path, "w", encoding="utf-8") as f:
            f.write(f"{self.doc_title.upper()}\n{'=' * len(self.doc_title)}\n\n")
            for sec in self.sections:
                if sec['type'] == 'heading':
                    f.write(f"\n# {sec['content']}\n")
                else:
                    f.write(f"{sec['content']}\n\n")
        exported_files.append("TXT")

        # 2. Markdown (.md)
        md_path = f"{base_path}.md"
        with open(md_path, "w", encoding="utf-8") as f:
            f.write(f"# {self.doc_title}\n\n")
            for sec in self.sections:
                if sec['type'] == 'heading':
                    f.write(f"## {sec['content']}\n\n")
                else:
                    f.write(f"{sec['content']}\n\n")
        exported_files.append("Markdown")

        # 3. HTML (.html)
        html_path = f"{base_path}.html"
        with open(md_path, "r", encoding="utf-8") as f:
            md_text = f.read()
        html_body = markdown.markdown(md_text)
        with open(html_path, "w", encoding="utf-8") as f:
            f.write(f"<!DOCTYPE html><html><head><title>{self.doc_title}</title></head><body>{html_body}</body></html>")
        exported_files.append("HTML")

        # 4. Word (.docx)
        docx_path = f"{base_path}.docx"
        doc = Document()
        doc.add_heading(self.doc_title, level=0)
        for sec in self.sections:
            if sec['type'] == 'heading':
                doc.add_heading(sec['content'], level=1)
            else:
                doc.add_paragraph(sec['content'])
        doc.save(docx_path)
        exported_files.append("Word DOCX")

        # 5. PDF (.pdf)
        pdf_path = f"{base_path}.pdf"
        pdf = FPDF()
        pdf.add_page()
        pdf.set_font("Helvetica", "B", 18)
        pdf.cell(0, 10, self.doc_title, new_x="LMARGIN", new_y="NEXT", align="C")
        pdf.ln(5)
        
        for sec in self.sections:
            if sec['type'] == 'heading':
                pdf.set_font("Helvetica", "B", 14)
                pdf.cell(0, 8, sec['content'], new_x="LMARGIN", new_y="NEXT")
            else:
                pdf.set_font("Helvetica", "", 11)
                pdf.multi_cell(0, 6, sec['content'])
                pdf.ln(3)
        pdf.output(pdf_path)
        exported_files.append("PDF")

        return f"Successfully exported '{self.doc_title}' to {', '.join(exported_files)} in {self.output_dir}."

    def share_to_dropbox_or_cloud(self, cloud_folder=None):
        """Share document to Dropbox/Nextcloud sync directory."""
        if not cloud_folder:
            cloud_folder = os.path.expanduser("~/Dropbox/ArtomeDocs")
        
        os.makedirs(cloud_folder, exist_ok=True)
        filename_base = self.doc_title.lower().replace(" ", "_")

        # Copy exported PDF and DOCX to cloud directory
        pdf_src = os.path.join(self.output_dir, f"{filename_base}.pdf")
        docx_src = os.path.join(self.output_dir, f"{filename_base}.docx")

        if not os.path.exists(pdf_src):
            self.export_all_formats(filename_base)

        synced = []
        if os.path.exists(pdf_src):
            shutil.copy(pdf_src, os.path.join(cloud_folder, f"{filename_base}.pdf"))
            synced.append("PDF")
        if os.path.exists(docx_src):
            shutil.copy(docx_src, os.path.join(cloud_folder, f"{filename_base}.docx"))
            synced.append("DOCX")

        return f"Document '{self.doc_title}' shared and synced to Dropbox/Cloud folder ({', '.join(synced)})."

if __name__ == "__main__":
    writer = AudioDocWriter()
    print(writer.start_new_doc("Artome Project Report"))
    print(writer.add_heading("Executive Summary"))
    print(writer.add_paragraph("Artome OS is a voice-first conversational desktop environment for the blind."))
    print(writer.add_heading("Features"))
    print(writer.add_paragraph("Includes Audio IDE, File Browser, Web Browser, Email Client, and Multi-Format Document Writer."))
    print("\nExporting all formats...")
    print(writer.export_all_formats())
